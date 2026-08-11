# -*- coding: utf-8 -*-
"""
生成 OmniAssist（轻量级 AI Agent 服务）四份文档：
  1. 功能清单.xlsx
  2. 产品说明.docx
  3. 需求设计文档.docx
  4. 接口文档.docx
所有内容均依据实际代码（server/、agent/、static/）梳理。
"""
import os
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "docs")
os.makedirs(OUT_DIR, exist_ok=True)

# =====================================================================
# 公共样式
# =====================================================================
ACCENT = RGBColor(0x1F, 0x6F, 0xB2)      # 主题蓝
ACCENT_HEX = "1F6FB2"
HDR_FILL = "1F6FB2"
ALT_FILL = "EAF2F8"


def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = ACCENT
    return h


def add_para(doc, text, bold=False, size=None, italic=False, space_after=6):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    if size:
        r.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 + 0.25 * level)
    r = p.add_run(text)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style='List Number')
    r = p.add_run(text)
    return p


def add_doc_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(9)
        set_cell_bg(hdr[i], HDR_FILL)
    for ri, row in enumerate(rows):
        cells = table.add_row().cells
        for ci, val in enumerate(row):
            cells[ci].text = ""
            run = cells[ci].paragraphs[0].add_run(str(val))
            run.font.size = Pt(9)
            if ri % 2 == 1:
                set_cell_bg(cells[ci], ALT_FILL)
    if col_widths:
        for ci, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[ci].width = Inches(w)
    return table


def add_code(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = 'Consolas'
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(4)
    return p


# =====================================================================
# 一、功能清单数据
# =====================================================================
# (模块, 功能名称, 功能描述, 关键能力/特性, 适用角色, 相关接口/模块)
FUNC_ROWS = [
    ("用户与认证", "账号登录", "用户使用用户名+密码登录平台，成功后发放 HMAC-SHA256 签名 Token（30分钟过期）。",
     "Token 认证、Cookie 持久化、首次登录强制改密", "全部角色", "/api/auth/login"),
    ("用户与认证", "账号登出", "清除当前会话 Token，结束登录态。", "无状态登出", "全部角色", "/api/auth/logout"),
    ("用户与认证", "修改密码", "用户修改自身密码，需校验旧密码且新密码≥6位，两次输入一致。", "强度校验、DB Web 代理密码同步", "全部角色", "/api/auth/password"),
    ("用户与认证", "当前用户信息", "获取当前登录用户的 id、用户名、角色类型与描述。", "身份透出", "全部角色", "/api/auth/me"),
    ("用户与认证", "权限查询", "返回当前用户所属角色及其权限矩阵。", "RBAC 可视化", "全部角色", "/api/auth/permissions"),
    ("用户与认证", "用户管理（增删改查）", "管理员可创建/编辑/删除用户并分配角色（admin/user）。", "角色分配、禁止自删、可选保留文件", "管理员", "/api/users"),
    ("用户与认证", "基于角色的访问控制", "路由级权限校验（users/model_config_global/search_config 等），越权返回 403。",
     "RBAC 九类资源权限", "全部角色", "server/database.py permissions"),

    ("会话管理", "创建会话", "新建对话会话；首次在聊天框发送消息自动创建新会话（未在任何会话中时）。",
     "标题长度≤200、默认标题“新对话”", "全部角色", "/api/sessions POST"),
    ("会话管理", "会话列表", "列出当前用户的所有会话（按时间）。", "按用户隔离", "全部角色", "/api/sessions GET"),
    ("会话管理", "会话搜索", "按关键词搜索当前用户的会话标题/内容。", "模糊匹配", "全部角色", "/api/sessions/search"),
    ("会话管理", "会话详情", "获取会话完整历史消息（含思考、工具调用、联网记录）。", "no-cache 响应头", "全部角色", "/api/sessions/{id} GET"),
    ("会话管理", "重命名会话", "修改会话标题，长度≤200字符。", "标题校验", "全部角色", "/api/sessions/{id} PUT"),
    ("会话管理", "删除会话", "删除会话并清理其信任模式状态与运行中任务。", "级联清理", "全部角色", "/api/sessions/{id} DELETE"),
    ("会话管理", "任务状态查询", "查询当前用户所有运行中/最近完成的后台任务状态。", "多任务监控", "全部角色", "/api/sessions/task-status/all"),
    ("会话管理", "顶部会话标题栏", "页面顶部展示当前会话标题（与侧边栏同源），不遮挡模式按钮。", "UX 优化", "全部角色", "static/index.html"),

    ("对话与推理引擎", "流式对话", "基于 SSE 的实时流式回答输出，回答随生成逐步呈现。", "text/event-stream、后端任务解耦", "全部角色", "/api/chat/stream"),
    ("对话与推理引擎", "事件订阅/重连", "切换会话后重新订阅正在运行的任务事件流。", "多订阅者缓冲回放", "全部角色", "/api/chat/subscribe/{id}"),
    ("对话与推理引擎", "中途停止", "停止当前会话正在运行的后台任务并取消待审请求。", "asyncio.Task.cancel", "全部角色", "/api/chat/stop/{id}"),
    ("对话与推理引擎", "联网搜索", "集成 Tavily 搜索引擎，智能识别 7 类场景（实时/事实/最新/教程/对比/本地化/通用）并优化策略。",
     "意图识别、可折叠结果面板", "全部角色", "agent/llm + Tavily"),
    ("对话与推理引擎", "思考过程可视化", "可切换显示/隐藏模型思考过程（自然独白），思考→联网→工具→回答四段式折叠展示。", "show_thought 开关", "全部角色", "/api/config PUT"),
    ("对话与推理引擎", "上下文压缩", "设置上下文 token 上限（32k/64k/128k 等），自动压缩历史消息避免超限。", "上下文管理", "全部角色", "agent/agent.py"),
    ("对话与推理引擎", "意图关键词自优化", "基于用户输入动态匹配意图关键词，按用户存储并自动优化工具选择。", "降 token 开销", "全部角色", "agent/intent_keywords.py"),
    ("对话与推理引擎", "多轮对话/历史还原", "Per-Session DOM 隔离，切换会话完整还原历史四段式布局。", "会话持久化", "全部角色", "server/database.py"),
    ("对话与推理引擎", "斜杠命令", "提供 /help、/reset 等简易命令。", "交互辅助", "全部角色", "/api/chat/commands"),
    ("对话与推理引擎", "输入安全校验", "单条消息≤10000字符、标题≤200字符、不存在会话返回 404。", "防资源耗尽", "全部角色", "server/routes/chat.py"),

    ("技能系统", "系统技能", "预置并维护 calculator、datetime、web-fetch、document（Word/Excel/PPT/PDF）、gold-price 等系统技能。",
     "文件系统仓库、管理员维护", "全部角色(用)/管理员(改)", "agent/skills/"),
    ("技能系统", "用户技能 CRUD", "创建/查看/更新/删除用户自定义技能（SKILL.md + 工具脚本）。", "数据库持久化、缓存清理", "管理员(写)/全员(读)", "/api/skills"),
    ("技能系统", "技能启停", "按用户隔离启用/禁用系统技能与用户技能（禁用标记机制）。", "按用户状态隔离", "全部角色", "/api/skills/{id}/toggle"),
    ("技能系统", "技能模板安装", "将官方技能模板（需自带 API Key，如 weather）安装到用户技能仓库。", "install_skill_template", "全部角色", "main.py install_skill_template"),
    ("技能系统", "自然语言创建技能", "对话中描述需求即可让 Agent 自动生成、验证并保存技能（admin 专属）。", "端到端创建", "管理员", "skill_editor.py"),
    ("技能系统", "技能注册中心", "系统/用户双仓库隔离，同名用户技能覆盖系统技能，统一加载为工具。", "SkillRegistry", "系统", "agent/skill_registry.py"),

    ("工具执行与沙箱", "三种执行模式", "local_execution（本地代码）/http_request（外部 API）/llm_simulated（模型直出）。", "模式路由", "系统", "agent/tools.py"),
    ("工具执行与沙箱", "工具沙箱隔离", "工具代码在独立 venv 与子进程执行，崩溃不影响主服务；每用户独立沙箱目录。", "SandboxPool、venv 隔离", "系统", "agent/sandbox.py"),
    ("工具执行与沙箱", "9 层安全防护", "进程隔离、venv 隔离、导入拦截、OS 危险函数禁用、文件删除控制、网络熔断（可配）、超时、环境隔离、用户隔离。",
     "纵深防御", "系统", "agent/sandbox.py"),
    ("工具执行与沙箱", "依赖管理", "工具按需 pip 安装依赖，共享基础环境经 PYTHONPATH 继承、按用户记录避免重复安装。", "镜像源可配", "系统", "agent/sandbox.py"),
    ("工具执行与沙箱", "条件并行执行", "多工具调用自动并行（线程池），单工具直接执行，避免不必要开销。", "ParallelToolExecutor", "系统", "agent/parallel_executor.py"),
    ("工具执行与沙箱", "子 Agent 调度", "多子 Agent 配置按需加载，AgentPool 统一管理；支持临时候时 Agent 处理一次性任务。", "AgentPool", "系统", "agent/agent_pool.py"),
    ("工具执行与沙箱", "受控命令执行", "沙箱内 run_command 仅允许白名单命令（python/node/ls/cat 等），默认禁用 rm/sudo。", "命令白名单", "系统", "agent/sandbox.py"),

    ("审批与权限模式", "请求批准模式", "敏感操作前在聊天框弹出审批卡片，逐项确认后才执行。", "request 模式", "全部角色", "/api/chat/{id}/trust"),
    ("审批与权限模式", "完全访问模式", "放开审批门（提示风险），适合高信任场景。", "full 模式、警告弹窗", "全部角色", "/api/chat/{id}/trust"),
    ("审批与权限模式", "三选项审批", "审批卡片提供“跳过/允许/拒绝”三项，点击即执行，支持增量逐项决议。", "skip/approve/reject", "全部角色", "/api/chat/{id}/approve"),
    ("审批与权限模式", "角色级免确认", "持有 tools:execute_sensitive 权限的角色可跳过敏感操作确认。", "trust_store", "管理员", "chat.py _is_role_exempt"),
    ("审批与权限模式", "信任模式状态", "按会话+用户记录权限模式，可查询与切换。", "会话级隔离", "全部角色", "/api/chat/{id}/trust GET/POST"),
    ("审批与权限模式", "审批审计", "每次审批决议写入 logs/audit.log，可追溯。", "审计留痕", "系统", "server/routes/approval.py"),

    ("文件管理", "统一文件库", "合并“生成文件”与“上传文件”，按类别（Word/Excel/PDF/PPT/CSV/图片/文本）展示，支持搜索与管理员跨用户查看。",
     "document_output/{uid}/", "全部角色", "/api/files/library"),
    ("文件管理", "文件上传", "按会话上传文件（≤5个、≤20MB、扩展名白名单），自动解析与摘要。", "类型识别、摘要生成", "全部角色", "/api/files/upload"),
    ("文件管理", "文件列表", "按用户隔离列出生成文件目录结构。", "目录树", "全部角色", "/api/files GET"),
    ("文件管理", "文件下载", "下载或内联预览文件，带访问权限校验。", "inline 参数", "全部角色", "/api/files/download"),
    ("文件管理", "文件预览", "支持文本/代码/图片/CSV(表格)/docx(纯文本)/xlsx(共享串)/PDF 等在线预览，不支持格式友好提示。",
     "多格式适配", "全部角色", "/api/files/preview"),
    ("文件管理", "文件重命名", "文件库内同目录重命名，字符白名单与访问校验。", "防路径穿越", "全部角色", "/api/files/rename"),
    ("文件管理", "文件删除", "删除生成/上传文件，越权拒绝。", "权限校验", "全部角色", "/api/files DELETE"),
    ("文件管理", "上传文件管理", "列出/删除当前会话或全部上传文件，支持关键词搜索。", "all-uploads", "全部角色", "/api/files/uploads、/all-uploads"),
    ("文件管理", "引用文件", "将已有上传文件引用到目标会话（复制），受数量上限约束。", "reference-files", "全部角色", "/api/files/reference-files"),

    ("模型与搜索配置", "个人模型配置", "用户配置个人 API Key、Base URL、Model Name、上下文上限、思考模式开关。", "配置持久化、即时生效", "全部角色", "/api/config PUT"),
    ("模型与搜索配置", "全局模型配置", "管理员配置全局默认模型参数（优先级低于用户个人配置）。", "admin only", "管理员", "/api/config/global"),
    ("模型与搜索配置", "搜索配置", "管理员配置 Tavily 搜索 API Key（脱敏展示）。", "search_config", "管理员", "/api/config/search"),
    ("模型与搜索配置", "多模型网关", "根据模型名自动适配思考参数（native/extra_body/always_on/prompt）与温度策略。",
     "Doubao/GPT-5/O/DeepSeek/Qwen/Claude 等", "系统", "agent/model_gateway.py"),

    ("用户密钥与凭证", "用户密钥保存", "保存用户私有密钥/配置（如第三方 API Key、私有 host），按用户加密隔离。", "set_user_secret", "全部角色", "main.py set_user_secret"),
    ("用户密钥与凭证", "密钥列表", "列出当前用户密钥名称及脱敏值（不返回明文）。", "脱敏展示", "全部角色", "main.py list_user_secrets"),
    ("用户密钥与凭证", "密钥删除", "删除用户私有密钥。", "delete_user_secret", "全部角色", "main.py delete_user_secret"),

    ("任务复盘", "任务执行回顾", "查看最近任务执行记录（成功/失败、工具使用、错误信息）。", "review_recent_tasks", "全部角色", "main.py review_recent_tasks"),
    ("任务复盘", "失败模式分析", "分析任务日志，自动发现失败模式并生成技能优化建议。", "analyze_and_suggest", "全部角色", "main.py analyze_and_suggest"),

    ("运维诊断（管理员）", "日志查看", "读取应用/错误日志尾部（只读，行数上限）。", "diag_read_logs", "管理员", "agent/diagnostics.py"),
    ("运维诊断（管理员）", "服务状态检查", "查询端口监听、进程 PID、Python/pip 版本。", "diag_service_status", "管理员", "agent/diagnostics.py"),
    ("运维诊断（管理员）", "受控重启", "重启主服务（当前连接短暂中断）。", "diag_restart_service", "管理员", "agent/diagnostics.py"),
    ("运维诊断（管理员） " , "环境检查", "结构化只读检查 python 版本/pip 包/文件存在/端口监听。", "diag_check_env", "管理员", "agent/diagnostics.py"),
    ("运维诊断（管理员）", "数据库只读查询", "白名单表 SELECT 查询（user_skills/sessions/users/permissions/model_configs/search_configs）。", "diag_db_query", "管理员", "agent/diagnostics.py"),
    ("运维诊断（管理员）", "文件/技能校验", "读取白名单目录文件、校验 SKILL.md 格式、列出/重命名/删除生成文件。", "diag_* 系列", "管理员", "agent/diagnostics.py"),

    ("安全机制", "Token 认证", "HMAC-SHA256 签名 Token，30 分钟过期，Cookie/Authorization 头双支持。", "auth_middleware", "系统", "server/routes/auth.py"),
    ("安全机制", "密码与密钥加密", "密码 SHA256 哈希存储；API Key 采用 XOR+SHA256 派生密钥加密。", "加密存储", "系统", "server/database.py、agent/config.py"),
    ("安全机制", "输入校验", "消息/标题长度限制、会话存在性校验、路径穿越防护。", "边界防护", "系统", "server/routes/*"),
    ("安全机制", "路径与文件安全", "文件操作限定项目内、密钥目录禁读、写仅白名单目录。", "沙箱文件边界", "系统", "agent/sandbox.py、files.py"),
    ("安全机制", "数据库管理代理", "sqlite-web 仅绑定 127.0.0.1，经 Basic 认证代理，避免公网暴露。", "本地回环", "管理员", "server/main.py _start_sqlite_web"),

    ("前端界面与体验", "双主题模式", "亮色/暗色/自动（跟随系统）三态循环切换。", "主题切换", "全部角色", "static/style.css"),
    ("前端界面与体验", "登录页美化", "毛玻璃登录框、浮动光斑背景、四时段渐变（早/午/晚/夜）。", "backdrop-filter", "全部角色", "static/login.html"),
    ("前端界面与体验", "输入工具栏", "输入框下方“联网/思考/请求批准”三按钮，与输入框留间距。", "模式切换", "全部角色", "static/index.html"),
    ("前端界面与体验", "审批卡片", "聊天框内审批卡片，三选项点击即执行，状态固化展示。", "UX 交互", "全部角色", "static/app.js"),
    ("前端界面与体验", "文件库界面", "操作按钮一行、来源徽标、多格式预览弹窗。", "UX 交互", "全部角色", "static/app.js"),
    ("前端界面与体验", "历史对话面板", "切换会话完整还原、折叠/展开、来源分类标签。", "UX 交互", "全部角色", "static/app.js"),
]


def gen_xlsx():
    wb = Workbook()
    ws = wb.active
    ws.title = "功能清单"
    headers = ["序号", "功能模块", "功能名称", "功能描述", "关键能力/特性", "适用角色", "相关接口/模块"]
    ws.append(headers)
    # 表头样式
    thin = Side(style="thin", color="BBBBBB")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    for c in range(1, len(headers) + 1):
        cell = ws.cell(row=1, column=c)
        cell.font = Font(bold=True, color="FFFFFF", size=11)
        cell.fill = PatternFill("solid", fgColor=HDR_FILL)
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cell.border = border
    # 数据
    for i, row in enumerate(FUNC_ROWS, start=1):
        ws.append([i] + list(row))
        r = i + 1
        for c in range(1, len(headers) + 1):
            cell = ws.cell(row=r, column=c)
            cell.alignment = Alignment(vertical="top", wrap_text=True)
            cell.border = border
            if r % 2 == 0:
                cell.fill = PatternFill("solid", fgColor=ALT_FILL)
            if c == 1:
                cell.alignment = Alignment(horizontal="center", vertical="top")
    widths = [6, 18, 22, 52, 32, 16, 26]
    for i, w in enumerate(widths, start=1):
        ws.column_dimensions[get_column_letter(i)].width = w
    ws.freeze_panes = "A2"
    ws.row_dimensions[1].height = 24
    # 统计 sheet
    ws2 = wb.create_sheet("模块统计")
    ws2.append(["功能模块", "功能项数量"])
    module_count = {}
    for row in FUNC_ROWS:
        module_count[row[0]] = module_count.get(row[0], 0) + 1
    for m, cnt in module_count.items():
        ws2.append([m, cnt])
    ws2.append(["合计", len(FUNC_ROWS)])
    for c in range(1, 3):
        cell = ws2.cell(row=1, column=c)
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor=HDR_FILL)
        cell.alignment = Alignment(horizontal="center")
    for i in range(2, ws2.max_row + 1):
        ws2.cell(row=i, column=1).border = border
        ws2.cell(row=i, column=2).border = border
    ws2.column_dimensions["A"].width = 22
    ws2.column_dimensions["B"].width = 14
    path = os.path.join(OUT_DIR, "功能清单.xlsx")
    wb.save(path)
    return path


# =====================================================================
# 二、产品说明.docx
# =====================================================================
def gen_product_docx():
    doc = Document()
    # 默认字体
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style.font.size = Pt(10.5)

    title = doc.add_heading('OmniAssist 产品说明', level=0)
    for r in title.runs:
        r.font.color.rgb = ACCENT
    add_para(doc, "版本：v1.0    |    文档类型：产品说明    |    适用对象：产品/研发/运营/管理员", italic=True, size=9)

    add_heading(doc, "1. 产品概述", 1)
    add_para(doc, "OmniAssist 是一个轻量级 AI Agent 服务框架，基于“技能系统（Skill）”实现工具调用，"
                  "支持以自然语言动态创建与管理技能，无需编写代码即可扩展 Agent 能力。"
                  "平台同时提供 Web 前端与终端命令行两种交互方式，共享同一套 Agent 核心。")
    add_para(doc, "核心价值：把“调用工具/写代码/对接外部 API”的能力封装为可复用、可审批、可隔离执行的技能；"
                  "在保障安全（沙箱隔离 + 分级审批）的前提下，让普通用户也能通过对话完成复杂任务。")
    add_bullet(doc, "定位：面向个人与团队的轻量 Agent 中台，强调“安全可控的工具执行”。")
    add_bullet(doc, "差异化：内建 9 层沙箱防护、请求批准/完全访问双权限模式、按用户隔离的 venv 执行环境。")
    add_bullet(doc, "目标用户：需要把重复工作自动化的知识工作者（普通用户）；负责技能沉淀与运维的管理员。")

    add_heading(doc, "2. 产品功能架构", 1)
    add_para(doc, "平台采用分层结构，自上而下分为交互层、服务层、核心层、安全层与数据层。")

    add_heading(doc, "2.1 交互层", 2)
    add_bullet(doc, "Web 前端（static/）：对话主界面、会话侧边栏、设置面板、用户/技能/文件管理、登录页。")
    add_bullet(doc, "终端命令行（agent/main.py）：面向开发者的本地交互入口，支持 /model set 等命令。")

    add_heading(doc, "2.2 服务层（FastAPI，server/routes/）", 2)
    add_doc_table(doc,
        ["服务模块", "职责", "主要接口前缀"],
        [
            ["认证 auth", "登录/登出/改密/Token/权限", "/api/auth"],
            ["会话 sessions", "会话增删改查/搜索/任务状态", "/api/sessions"],
            ["对话 chat", "流式对话/订阅/停止/审批/信任模式", "/api/chat"],
            ["审批 approval", "（并入 chat）审批决议与权限模式", "/api/chat/{id}/approve、/trust"],
            ["技能 skills", "技能 CRUD/启停/详情", "/api/skills"],
            ["文件 files", "文件库/预览/下载/重命名/删除", "/api/files"],
            ["上传 upload", "文件上传/引用/管理", "/api/files/upload*"],
            ["用户 users", "用户管理（管理员）", "/api/users"],
            ["配置 config", "模型/搜索配置", "/api/config"],
        ],
        col_widths=[1.4, 3.2, 2.0])

    add_heading(doc, "2.3 核心层（agent/）", 2)
    add_doc_table(doc,
        ["核心模块", "职责"],
        [
            ["Agent 核心 agent.py", "多轮对话、上下文压缩、工具编排调度"],
            ["技能注册 SkillRegistry", "系统/用户双仓库、加载、同名覆盖"],
            ["工具注册 ToolRegistry", "工具注册、风险分级、审批门封装"],
            ["模型网关 ModelGateway", "多模型思考/温度参数适配"],
            ["沙箱 SandboxPool", "每用户独立 venv、9 层防护、依赖管理"],
            ["子 Agent 池 AgentPool", "多子 Agent 加载与临时候时 Agent"],
            ["意图关键词 intent_keywords", "按需工具选择、降 token 开销"],
            ["任务复盘 task_reviewer", "执行日志、失败分析、优化建议"],
            ["文件解析 file_parser", "上传文件解析与摘要"],
        ],
        col_widths=[2.2, 4.4])

    add_heading(doc, "2.4 数据层与存储", 2)
    add_bullet(doc, "SQLite（WAL 模式）持久化：用户、会话消息、模型/搜索配置、权限、用户技能、审批审计。")
    add_bullet(doc, "文件存储：document_output/{user_id}/ 下按类别（word_output/excel_output/pdf_output/ppt_output/csv_output/image_output/text_output/uploads）存放生成与上传文件。")
    add_bullet(doc, "sqlite-web 数据库管理代理：仅绑定 127.0.0.1，经 Basic 认证代理，供管理员本地查看。")

    add_heading(doc, "2.5 安全层", 2)
    add_bullet(doc, "认证：HMAC-SHA256 Token（30 分钟过期），中间件统一拦截未登录请求。")
    add_bullet(doc, "授权：RBAC 角色权限（admin/user），路由级校验。")
    add_bullet(doc, "执行安全：9 层沙箱防护，详见第 4 章。")
    add_bullet(doc, "审批：请求批准模式下的逐项确认（跳过/允许/拒绝）。")

    add_heading(doc, "3. 技术架构", 1)
    add_bullet(doc, "后端：FastAPI + Uvicorn，原生 JS SSE 流式推送；asyncio.Future 实现审批门阻塞/唤醒。")
    add_bullet(doc, "前端：原生 HTML/CSS/JS（无构建步骤），主题支持亮/暗/自动，登录页毛玻璃与四时段背景。")
    add_bullet(doc, "模型：OpenAI 兼容 API（Any OpenAI-format endpoint），经 ModelGateway 适配多厂商思考参数。")
    add_bullet(doc, "执行：每用户独立 venv 子进程，512MB 内存上限，超时保护，依赖经 PYTHONPATH 继承共享环境。")
    add_bullet(doc, "部署：单进程服务（默认 17520）；数据库代理 17521（本地回环）；Python 3.10~3.14 兼容。")
    add_para(doc, "数据流示例（一次带工具调用的对话）：", bold=True, space_after=2)
    add_code(doc,
        "用户输入 → auth 中间件校验 Token → chat/stream 创建 SessionTask(后台 asyncio 任务)\n"
        "  → Agent 核心调用 LLM → 命中技能/工具 → ToolRegistry 封装\n"
        "  → 若为敏感操作且处于“请求批准”模式 → 阻塞于 ApprovalStore 的 Future\n"
        "  → 前端弹审批卡片 → 用户点击(跳过/允许/拒绝) → /approve 写入决议并唤醒 Future\n"
        "  → 沙箱执行工具(子进程/venv) → 结果回流 → SSE 推送给订阅的前端")

    add_heading(doc, "4. 各模块详细说明", 1)
    modules = [
        ("server/main.py", "应用入口：全局服务初始化（配置、LLM、工具注册、技能加载、子 Agent、诊断工具）、"
                            "CORS、认证与异常中间件、静态托管；启动 sqlite-web 本地代理；端口冲突自愈。"),
        ("server/routes/auth.py", "基于 HMAC-SHA256 的 Token 签发/校验；登录、登出、改密、当前用户、权限查询；"
                                  "require_permission 路由级鉴权。"),
        ("server/routes/chat.py", "对话核心：SessionTask 后台任务管理（事件缓冲+多订阅者）、SSE 流式输出、"
                                  "订阅/停止、意图关键词选工具、审批门协调、工具参数脱敏、/commands。"),
        ("server/routes/approval.py", "审批决议端点（增量逐项提交、会话归属校验、防越权）、信任/权限模式查询与切换、审计落盘。"),
        ("server/routes/sessions.py", "会话增删改查、搜索、任务状态查询；删除时级联清理信任状态与运行中任务。"),
        ("server/routes/skills.py", "技能列表/详情、用户技能 CRUD、启停（系统/数据库/文件系统三类处理）、禁用标记机制。"),
        ("server/routes/files.py", "统一文件库递归扫描、下载、预览（docx/csv/xlsx/pdf/图片/文本多格式）、重命名、删除。"),
        ("server/routes/upload.py", "文件上传（数量/大小/扩展名限制）、会话/全部上传列表、删除、跨会话引用。"),
        ("server/routes/users.py", "用户增删改查（管理员），禁止自删，可选保留文件。"),
        ("server/routes/config.py", "个人/全局模型配置、搜索配置（Tavily），Key 脱敏展示，即时热更新 Agent。"),
        ("agent/agent.py", "SimpleAgent：多轮对话、上下文压缩、工具编排。"),
        ("agent/sandbox.py", "ToolSandbox + SandboxPool：独立 venv、9 层防护（模块黑名单、OS 危险函数删除、"
                             "文件读禁区/写白名单、受控 run_command、内存/超时、网络熔断可配）、用户路径重写、执行日志。"),
        ("agent/tools.py", "ToolRegistry：工具注册、风险分级（read/exec/write/safe）、审批门封装、参数 Schema。"),
        ("agent/model_gateway.py", "MODEL_CAPABILITIES：按模型名（Doubao/GPT-5/O/DeepSeek/Qwen/Claude 等）"
                                   "映射思考参数类型与推理字段，归一化流式推理内容。"),
        ("agent/skill_registry.py", "SkillRegistry：系统技能（文件系统）与用户技能（数据库/文件系统）双仓库，加载、缓存、同名覆盖。"),
        ("agent/agent_pool.py", "多子 Agent 配置加载与统一调度，支持动态临时 Agent。"),
        ("agent/intent_keywords.py", "按需工具选择的关键词匹配与按用户自优化。"),
        ("agent/task_reviewer.py", "任务执行日志、失败模式分析、技能优化建议。"),
        ("agent/diagnostics.py", "管理员诊断工具集（日志/状态/重启/环境/DB 查询/文件），均带权限与白名单。"),
        ("agent/user_secrets.py / tool_secrets.py", "用户级密钥按用户加密隔离存取；工具运行时密钥注入沙箱。"),
        ("static/", "前端：index.html 主界面、login.html 登录、app.js 逻辑、style.css 主题与组件样式。"),
    ]
    for name, desc in modules:
        p = doc.add_paragraph()
        r = p.add_run(name)
        r.bold = True
        r.font.color.rgb = ACCENT
        p.add_run(" — " + desc)
        p.paragraph_format.space_after = Pt(4)

    add_heading(doc, "5. 核心机制", 1)
    add_heading(doc, "5.1 审批门（Approval Gate）", 2)
    add_para(doc, "当工具被判定为敏感（用户自建技能：exec/write 且 require_approval=True）且会话处于“请求批准”模式时，"
                  "执行流程在 asyncio.Future 上阻塞；前端弹出审批卡片，用户选择“跳过/允许/拒绝”后，"
                  "/approve 增量写入决议，所有项齐备后唤醒 Future 继续执行。系统技能默认免审批。"
                  "角色持有 tools:execute_sensitive 权限可跳过敏感确认（信任模式）。")
    add_heading(doc, "5.2 双权限模式", 2)
    add_bullet(doc, "请求批准（request）：敏感操作前逐项确认，安全优先。")
    add_bullet(doc, "完全访问（full）：放开审批门（含风险提示弹窗），效率优先。")
    add_bullet(doc, "状态按 会话+用户 维度隔离存储于 TrustStore。")
    add_heading(doc, "5.3 沙箱隔离", 2)
    add_bullet(doc, "每用户独立 venv 与子进程，依赖完全隔离；共享基础环境经 PYTHONPATH 只读继承避免重复安装。")
    add_bullet(doc, "9 层防护：进程隔离 / venv 隔离 / 导入拦截 / OS 危险函数禁用 / 文件删除控制 / 网络熔断(可配) / 超时 / 环境隔离 / 用户隔离。")
    add_bullet(doc, "内存上限 512MB（resource.setrlimit），超时强制终止，所有依赖安装真实错误写入日志。")
    add_heading(doc, "5.4 技能系统", 2)
    add_para(doc, "技能 = SKILL.md（YAML frontmatter + Markdown 指令）+ scripts/（Python 脚本）。"
                  "系统技能由管理员维护于 agent/skills/，用户技能存于 agent/skills/user/{user_id}/ 或数据库，"
                  "同名用户技能覆盖系统技能。加载后统一注册为工具，由 LLM 按需调用。")

    add_heading(doc, "6. 角色与权限矩阵", 1)
    add_doc_table(doc,
        ["角色", "核心权限"],
        [
            ["admin", "全部：用户管理、全局模型配置、搜索配置、技能读写、诊断工具、完全访问模式"],
            ["user", "基础：对话、会话管理、个人模型配置、技能查看、文件管理、请求批准模式"],
        ],
        col_widths=[1.4, 5.2])

    add_heading(doc, "7. 运行环境要求", 1)
    add_bullet(doc, "Python 3.10 及以上（兼容 3.10~3.14）。")
    add_bullet(doc, "OpenAI 兼容的模型 API 端点（任意兼容服务均可）。")
    add_bullet(doc, "（可选）Tavily API Key，用于联网搜索。")
    add_bullet(doc, "默认服务地址：http://localhost:17520；数据库代理：http://localhost:17521（仅本地）。")

    path = os.path.join(OUT_DIR, "产品说明.docx")
    doc.save(path)
    return path


# =====================================================================
# 三、需求设计文档.docx
# =====================================================================
def gen_requirements_docx():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style.font.size = Pt(10.5)

    title = doc.add_heading('OmniAssist 平台需求设计文档', level=0)
    for r in title.runs:
        r.font.color.rgb = ACCENT
    add_para(doc, "版本：v1.0    |    文档类型：需求设计    |    状态：已实现（本文档对照当前代码梳理）", italic=True, size=9)

    add_heading(doc, "1. 引言", 1)
    add_heading(doc, "1.1 编写目的", 2)
    add_para(doc, "本文档定义 OmniAssist 轻量级 AI Agent 服务平台的功能需求、非功能需求、关键流程与数据/安全设计，"
                  "作为研发实现、测试验证与后续迭代的依据。")
    add_heading(doc, "1.2 背景", 2)
    add_para(doc, "用户希望用自然语言驱动 Agent 完成“计算、生成文档、查询外部信息、执行代码”等任务，但需在不信任"
                  "自定义代码的前提下保障安全。平台以“技能系统 + 沙箱隔离 + 分级审批”解决该矛盾。")
    add_heading(doc, "1.3 术语", 2)
    add_bullet(doc, "Skill（技能）：可复用的能力单元（SKILL.md + 脚本），加载后成为工具。")
    add_bullet(doc, "Tool（工具）：注册到 Agent 的可调用能力，含风险分级。")
    add_bullet(doc, "Sandbox（沙箱）：隔离的工具执行环境（venv + 子进程）。")
    add_bullet(doc, "Approval Gate（审批门）：敏感操作前的用户确认机制。")
    add_bullet(doc, "RBAC：基于角色的访问控制。")

    add_heading(doc, "2. 总体描述", 1)
    add_heading(doc, "2.1 产品定位", 2)
    add_para(doc, "面向个人与团队的轻量 Agent 中台，强调“安全可控的工具执行”与“自然语言可扩展的能力”。")
    add_heading(doc, "2.2 用户特征", 2)
    add_bullet(doc, "普通用户：通过对话完成工作，关注易用性与结果可见性。")
    add_bullet(doc, "管理员：负责用户/模型/技能/运维，关注安全性、可控性与可观测性。")
    add_heading(doc, "2.3 运行环境", 2)
    add_bullet(doc, "服务端：FastAPI + Uvicorn，单进程；SQLite（WAL）存储；默认端口 17520。")
    add_bullet(doc, "客户端：现代浏览器（Chrome/Edge/Firefox/Safari），支持亮/暗/自动主题。")
    add_bullet(doc, "依赖：Python 3.10+，OpenAI 兼容模型 API，可选 Tavily。")
    add_heading(doc, "2.4 约束", 2)
    add_bullet(doc, "安全约束：用户自定义代码必须沙箱隔离、敏感操作可审批、密钥不落明文。")
    add_bullet(doc, "合规约束：管理员操作（用户/全局配置/诊断）需权限校验并留痕。")

    add_heading(doc, "3. 功能需求", 1)
    add_para(doc, "以下按模块列出功能需求，优先级：P0（核心必需）/ P1（重要）/ P2（增强）。完整功能项见《功能清单.xlsx》。")

    fr_modules = [
        ("FR-A 用户与认证", [
            ("FR-A1", "P0", "用户以 用户名+密码 登录，成功后获得 30 分钟有效 Token。"),
            ("FR-A2", "P0", "首次登录强制修改密码后方可使用平台。"),
            ("FR-A3", "P0", "用户可修改自身密码（≥6 位、两次一致）。"),
            ("FR-A4", "P0", "管理员可创建/编辑/删除用户并分配角色；禁止删除自身账号。"),
            ("FR-A5", "P0", "路由级 RBAC 校验，越权返回 403。"),
        ]),
        ("FR-B 会话管理", [
            ("FR-B1", "P0", "支持多会话创建、列表、搜索、详情、重命名、删除。"),
            ("FR-B2", "P0", "未在任何会话中时，聊天框发送消息自动创建新会话。"),
            ("FR-B3", "P0", "顶部展示当前会话标题；切换/删除会话级联清理运行状态。"),
            ("FR-B4", "P1", "可查询当前用户所有运行中/最近完成任务状态。"),
        ]),
        ("FR-C 对话与推理", [
            ("FR-C1", "P0", "基于 SSE 的流式回答输出，支持中途停止。"),
            ("FR-C2", "P0", "切换会话后可重连订阅运行中任务事件。"),
            ("FR-C3", "P0", "集成 Tavily 联网搜索，支持 7 类场景识别与结果折叠展示。"),
            ("FR-C4", "P1", "思考过程可视化（可开关），四段式（思考→联网→工具→回答）展示。"),
            ("FR-C5", "P1", "上下文 token 上限设置与自动压缩。"),
            ("FR-C6", "P2", "意图关键词自优化，提升工具选择精准度、降低 token 开销。"),
        ]),
        ("FR-D 技能系统", [
            ("FR-D1", "P0", "预置系统技能（计算/日期/文档生成/web-fetch/金价等）。"),
            ("FR-D2", "P0", "用户技能 CRUD（管理员），按用户隔离启用/禁用。"),
            ("FR-D3", "P1", "技能模板安装（需自带 API Key 的模板，如 weather）。"),
            ("FR-D4", "P1", "自然语言创建/更新/删除技能（管理员），端到端生成与自测。"),
            ("FR-D5", "P0", "同名用户技能覆盖系统技能，注册中心统一加载为工具。"),
        ]),
        ("FR-E 工具执行与沙箱", [
            ("FR-E1", "P0", "三种执行模式：local_execution / http_request / llm_simulated。"),
            ("FR-E2", "P0", "每用户独立 venv 子进程沙箱，崩溃不影响主服务。"),
            ("FR-E3", "P0", "9 层安全防护（模块黑名单、OS 危险函数禁用、文件读禁区/写白名单、受控命令、超时、内存上限、网络熔断可配等）。"),
            ("FR-E4", "P1", "多工具条件并行执行；子 Agent 调度。"),
        ]),
        ("FR-F 审批与权限模式", [
            ("FR-F1", "P0", "请求批准模式：敏感操作前弹审批卡片逐项确认。"),
            ("FR-F2", "P0", "完全访问模式：放开审批门（含风险提示）。"),
            ("FR-F3", "P0", "审批三选项（跳过/允许/拒绝），点击即执行，增量逐项决议。"),
            ("FR-F4", "P1", "角色级免确认（tools:execute_sensitive）。"),
            ("FR-F5", "P0", "每次审批写入审计日志，可追溯。"),
        ]),
        ("FR-G 文件管理", [
            ("FR-G1", "P0", "统一文件库（生成+上传），按类别展示，支持搜索。"),
            ("FR-G2", "P0", "文件上传（≤5 个、≤20MB、扩展名白名单）与解析摘要。"),
            ("FR-G3", "P0", "文件预览（文本/代码/图片/CSV/docx/xlsx/PDF 等）、下载、重命名、删除。"),
            ("FR-G4", "P1", "管理员可跨用户查看文件库；引用文件到目标会话。"),
        ]),
        ("FR-H 模型与配置", [
            ("FR-H1", "P0", "个人模型配置（API Key/Base URL/Model/上下文上限/思考开关）。"),
            ("FR-H2", "P0", "管理员全局模型配置与 Tavily 搜索配置（脱敏）。"),
            ("FR-H3", "P1", "多模型网关自动适配思考/温度参数。"),
        ]),
        ("FR-I 密钥与复盘/运维", [
            ("FR-I1", "P1", "用户密钥按用户加密隔离保存/列出(脱敏)/删除。"),
            ("FR-I2", "P2", "任务复盘：执行回顾与失败模式分析建议。"),
            ("FR-I3", "P1", "管理员诊断工具（日志/状态/重启/环境/DB 只读/文件），带权限与白名单。"),
        ]),
    ]
    for mname, items in fr_modules:
        add_heading(doc, mname, 2)
        add_doc_table(doc, ["需求编号", "优先级", "需求描述"], items, col_widths=[1.1, 0.8, 4.7])

    add_heading(doc, "4. 非功能需求", 1)
    add_doc_table(doc,
        ["类别", "需求"],
        [
            ["性能", "流式首字延迟低；后台任务与客户端连接解耦，支持多订阅者重连；工具执行超时保护。"],
            ["安全", "HMAC Token 认证；密码/密钥加密存储；9 层沙箱；RBAC；审批门；路径穿越与密钥禁读防护。"],
            ["可用性", "服务端口冲突自愈；沙箱初始化失败不影响对话；数据库代理仅本地回环。"],
            ["兼容性", "Python 3.10~3.14；OpenAI 兼容 API；现代浏览器；亮/暗/自动主题。"],
            ["可维护性", "分层架构（交互/服务/核心/安全/数据）；结构化日志按天轮转、错误分离；技能/工具可插拔。"],
            ["可观测性", "统一结构化日志、每用户沙箱执行日志、审批审计日志、诊断工具集。"],
            ["数据完整", "会话删除级联清理信任状态与运行中任务；SQLite WAL 模式。"],
        ],
        col_widths=[1.2, 5.4])

    add_heading(doc, "5. 关键流程设计", 1)
    add_heading(doc, "5.1 对话与工具调用流程", 2)
    add_number(doc, "前端 POST /api/chat/stream，服务端创建 SessionTask 后台 asyncio 任务并返回 SSE 流。")
    add_number(doc, "Agent 调用 LLM，若命中技能/工具则进入 ToolRegistry 封装。")
    add_number(doc, "若工具敏感且处于“请求批准”模式 → 阻塞于 ApprovalStore 的 Future，前端弹审批卡片。")
    add_number(doc, "用户点击 跳过/允许/拒绝 → /api/chat/{id}/approve 增量写入决议并唤醒 Future。")
    add_number(doc, "沙箱子进程执行工具 → 结果回流 → SSE 逐步推送给前端。")
    add_heading(doc, "5.2 技能创建流程（管理员）", 2)
    add_number(doc, "对话中描述需求 → Agent 理解并生成 SKILL.md 与脚本。")
    add_number(doc, "自动执行自测验证技能可用性。")
    add_number(doc, "保存至用户技能仓库（数据库/文件系统），清理缓存后立即可用。")
    add_heading(doc, "5.3 文件预览流程", 2)
    add_number(doc, "前端请求 /api/files/preview?path=...，服务端校验路径与权限。")
    add_number(doc, "按扩展名选择解析器（文本/CSV 表格/docx 文本/xlsx 串/图片/PDF 等）。")
    add_number(doc, "返回结构化预览数据，前端渲染弹窗。")

    add_heading(doc, "6. 数据需求", 1)
    add_para(doc, "采用 SQLite（WAL）。关键实体：users（用户与角色）、sessions（会话与消息 JSON）、"
                  "model_configs / search_configs（配置）、permissions（角色权限）、user_skills（用户技能与禁用标记）、"
                  "audit.log（审批审计）。文件实体以 document_output/{user_id}/ 目录树形式持久化。")
    add_doc_table(doc,
        ["数据实体", "说明"],
        [
            ["users", "id、username、password_hash、user_type(admin/user)、description、must_change_password"],
            ["sessions", "id、user_id、title、messages(JSON)、created_at"],
            ["model_configs", "user_id(NULL=全局)、api_key(加密)、base_url、model_name、context_limit、show_thought"],
            ["search_configs", "tavily_api_key(加密)"],
            ["permissions", "role、resource、action（RBAC 矩阵）"],
            ["user_skills", "id、user_id、skill_name、skill_content、skill_scripts、enabled（禁用标记复用）"],
        ],
        col_widths=[1.6, 5.0])

    add_heading(doc, "7. 安全与权限设计", 1)
    add_bullet(doc, "认证：HMAC-SHA256 Token（30 分钟），中间件统一拦截；Cookie/Authorization 双来源。")
    add_bullet(doc, "授权：RBAC（admin/user），路由级 require_permission 校验。")
    add_bullet(doc, "执行安全：9 层沙箱；用户自建技能标 exec/write 且 require_approval=True 必过审批门；系统技能免审批。")
    add_bullet(doc, "密钥：密码 SHA256 哈希；API Key 采用 XOR+SHA256 派生密钥加密；沙箱禁读密钥目录。")
    add_bullet(doc, "审计：审批决议、权限模式切换、诊断操作均落日志，可追溯、防越权。")

    add_heading(doc, "8. 待办与后续规划", 1)
    add_bullet(doc, "多 worker 部署时 ApprovalStore/TrustStore 需迁移至 Redis（当前为内存态）。")
    add_bullet(doc, "增强更多文件格式预览（如 PPT 在线预览）。")
    add_bullet(doc, "技能市场/共享与版本管理。")
    add_bullet(doc, "更细粒度的权限与配额管理。")

    path = os.path.join(OUT_DIR, "需求设计文档.docx")
    doc.save(path)
    return path


# =====================================================================
# 四、接口文档.docx
# =====================================================================
def gen_api_docx():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style.font.size = Pt(10.5)

    title = doc.add_heading('OmniAssist 平台接口文档', level=0)
    for r in title.runs:
        r.font.color.rgb = ACCENT
    add_para(doc, "版本：v1.0    |    基础路径：http://localhost:17520    |    API 前缀：/api", italic=True, size=9)

    add_heading(doc, "1. 通用说明", 1)
    add_bullet(doc, "认证：除 /api/auth/login、/api/health、/login.html、/favicon.ico 及 /static/* 外，所有请求需携带有效 Token"
                    "（Cookie: auth_token 或 Header: Authorization: Bearer <token>）。未登录访问 API 返回 401。")
    add_bullet(doc, "授权：管理员接口（用户管理、全局/搜索配置、诊断）需对应 RBAC 权限，越权返回 403。")
    add_bullet(doc, "流式接口：/api/chat/stream 与 /api/chat/subscribe/{id} 返回 text/event-stream（SSE）。")
    add_bullet(doc, "错误码：401 未登录 / 403 越权 / 404 不存在 / 400 参数错误 / 409 冲突（如任务进行中）/ 413 消息过长 / 500 内部错误。")
    add_para(doc, "以下接口均以方法 + 路径列出，请求体为 JSON（标注 Query 的为查询参数）。", space_after=2)

    # 每张表：方法, 路径, 说明, 关键参数/响应, 权限
    api_groups = [
        ("认证 /auth", [
            ("POST", "/api/auth/login", "账号登录", "body: {username,password} → {token,user_type,username,must_change_password}", "公开"),
            ("POST", "/api/auth/logout", "登出", "—", "登录"),
            ("PUT", "/api/auth/password", "修改密码", "body: {old_password,new_password,confirm_password}", "登录"),
            ("GET", "/api/auth/me", "当前用户信息", "→ {id,username,user_type,description}", "登录"),
            ("GET", "/api/auth/permissions", "当前权限矩阵", "→ {role,permissions}", "登录"),
        ]),
        ("会话 /sessions", [
            ("GET", "/api/sessions", "会话列表", "→ 当前用户会话数组", "登录"),
            ("GET", "/api/sessions/search?q=", "搜索会话", "Query q(≥1) → 会话数组", "登录"),
            ("POST", "/api/sessions", "创建会话", "body:{title?} → {id,title,created_at}", "登录"),
            ("GET", "/api/sessions/{id}", "会话详情(含消息)", "→ {id,title,messages,...}", "属主"),
            ("PUT", "/api/sessions/{id}", "重命名会话", "body:{title}(≤200)", "属主"),
            ("DELETE", "/api/sessions/{id}", "删除会话", "级联清理信任/任务 → {success}", "属主"),
            ("GET", "/api/sessions/task-status/all", "全部任务状态", "→ 任务状态映射", "登录"),
        ]),
        ("对话 /chat", [
            ("POST", "/api/chat/stream", "流式对话", "body:{session_id,message,web_search,show_thought} → SSE", "登录"),
            ("GET", "/api/chat/subscribe/{id}", "订阅任务流", "→ SSE（重连）", "属主"),
            ("POST", "/api/chat/stop/{id}", "停止任务", "→ {success}", "属主"),
            ("GET", "/api/chat/commands", "斜杠命令列表", "→ [{command,description,category}]", "登录"),
            ("POST", "/api/chat/{id}/approve", "提交审批决议", "body:{group_id,decisions:[{item_id,decision}]} decision∈{approve,reject,skip}", "属主"),
            ("GET", "/api/chat/{id}/trust", "查询权限模式", "→ {enabled,mode} mode∈{request,full}", "属主/管理员"),
            ("POST", "/api/chat/{id}/trust", "切换权限模式", "body:{enabled?,mode?} mode∈{request,full}", "属主/管理员"),
        ]),
        ("技能 /skills", [
            ("GET", "/api/skills?user_only=", "技能列表", "→ {system_skills,user_skills}", "登录"),
            ("GET", "/api/skills/system/{name}", "系统技能详情", "→ {name,description,instructions,scripts}", "登录"),
            ("GET", "/api/skills/user/{name}", "用户技能详情(按名)", "→ 详情", "属主"),
            ("GET", "/api/skills/{id}", "技能详情(按id)", "→ 详情", "属主"),
            ("POST", "/api/skills", "创建用户技能", "body:{name,content,scripts}", "管理员(写)"),
            ("PUT", "/api/skills/{id}", "更新用户技能", "body:{content?,scripts?,enabled?}", "管理员(写)"),
            ("DELETE", "/api/skills/{id}", "删除用户技能", "→ {message}", "管理员(写)"),
            ("PUT", "/api/skills/{id}/toggle", "启用/禁用技能", "body:{name,enabled}", "登录"),
        ]),
        ("文件 /files", [
            ("GET", "/api/files", "生成文件目录树", "→ 按用户隔离的文件树", "登录"),
            ("GET", "/api/files/library?search=&user_id=", "统一文件库", "→ {files,count,is_admin,target_uid}", "登录(管理员可跨用户)"),
            ("POST", "/api/files/rename", "重命名文件", "body:{path,new_name}", "属主"),
            ("GET", "/api/files/download?path=&inline=", "下载/内联预览", "→ 文件流", "属主"),
            ("DELETE", "/api/files?path=", "删除文件", "→ {success}", "属主"),
            ("GET", "/api/files/preview?path=", "在线预览", "→ {type,content/html/path,...}", "属主"),
        ]),
        ("上传 /files", [
            ("POST", "/api/files/upload?session_id=", "上传文件", "multipart files(≤5,≤20MB) → {uploaded,errors}", "登录"),
            ("GET", "/api/files/uploads?session_id=", "会话上传列表", "→ {files,count}", "属主"),
            ("DELETE", "/api/files/upload?session_id=&filename=", "删除上传", "→ {success}", "属主"),
            ("GET", "/api/files/all-uploads?search=", "全部上传列表", "→ {files,count}", "登录"),
            ("DELETE", "/api/files/all-uploads?path=", "删除指定上传", "→ {success}", "属主"),
            ("POST", "/api/files/reference-files", "引用文件到会话", "body:{paths,session_id}", "登录"),
        ]),
        ("用户 /users（管理员）", [
            ("GET", "/api/users", "用户列表", "→ UserResponse[]", "users:read"),
            ("POST", "/api/users", "创建用户", "body:{username,password,user_type,description}", "users:write"),
            ("PUT", "/api/users/{id}", "更新用户", "body:{password?,user_type?,description?}", "users:write"),
            ("DELETE", "/api/users/{id}?keep_files=", "删除用户", "禁止自删", "users:delete"),
        ]),
        ("配置 /config", [
            ("GET", "/api/config", "个人模型配置", "→ {model_name,base_url,api_key_masked,context_limit,show_thought,config_type}", "登录"),
            ("PUT", "/api/config", "更新个人配置", "body:{api_key?,base_url?,model_name?,context_limit?,show_thought?}", "登录"),
            ("GET", "/api/config/global", "全局模型配置", "→ 配置对象", "model_config_global:read"),
            ("PUT", "/api/config/global", "更新全局配置", "body 同个人", "model_config_global:write"),
            ("GET", "/api/config/search", "搜索配置", "→ {tavily_api_key_masked}", "search_config:read"),
            ("PUT", "/api/config/search", "更新搜索配置", "body:{tavily_api_key}", "search_config:write"),
        ]),
        ("系统", [
            ("GET", "/api/health", "健康检查", "→ {status:'ok'}", "公开"),
            ("GET", "/", "主页面", "→ index.html", "登录"),
            ("GET", "/login.html", "登录页", "→ login.html", "公开"),
        ]),
    ]

    for gname, rows in api_groups:
        add_heading(doc, gname, 2)
        add_doc_table(doc,
            ["方法", "路径", "说明", "请求/响应要点", "权限"],
            rows,
            col_widths=[0.7, 2.0, 1.2, 2.4, 1.0])

    add_heading(doc, "2. 审批决议示例", 1)
    add_para(doc, "请求批准模式下，前端在用户点击后调用：", space_after=2)
    add_code(doc,
        "POST /api/chat/{session_id}/approve\n"
        "{\n"
        '  "group_id": "grp_xxx",\n'
        '  "decisions": [\n'
        '    {"item_id": "item_1", "decision": "approve"},\n'
        '    {"item_id": "item_2", "decision": "skip"}\n'
        "  ]\n"
        "}\n"
        "→ 200 {\"success\": true}")
    add_para(doc, "decision 取值：approve（允许）/ reject（拒绝）/ skip（跳过，不执行）。"
                  "所有项决议齐备后唤醒被阻塞的工具执行。", size=9, italic=True)

    add_heading(doc, "3. 权限模式切换示例", 1)
    add_code(doc,
        "POST /api/chat/{session_id}/trust\n"
        '{"mode": "full"}        # 完全访问\n'
        '{"mode": "request"}     # 请求批准\n'
        "→ 200 {\"success\": true, \"enabled\": true, \"mode\": \"full\"}")

    path = os.path.join(OUT_DIR, "接口文档.docx")
    doc.save(path)
    return path


if __name__ == "__main__":
    p1 = gen_xlsx()
    print("生成:", p1)
    p2 = gen_product_docx()
    print("生成:", p2)
    p3 = gen_requirements_docx()
    print("生成:", p3)
    p4 = gen_api_docx()
    print("生成:", p4)
    print("全部完成。")

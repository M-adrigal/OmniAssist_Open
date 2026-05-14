"""
Word 文档格式化引擎

将结构化的排版描述（JSON）转换为 python-docx 文档。
支持字体、段落、页面、表格、列表、图片等全部排版维度。

设计原则：
- 所有尺寸单位统一使用厘米(cm)或磅(pt)，内部自动转换
- 未指定的属性使用合理默认值，不强制要求完整描述
- 无法精确实现的样式自动降级并在结果中提示
"""

import os

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

PAPER_SIZES = {
    "A4": (21.0, 29.7),
    "A3": (29.7, 42.0),
    "A5": (14.8, 21.0),
    "B5": (17.6, 25.0),
    "Letter": (21.59, 27.94),
    "Legal": (21.59, 35.56),
}

ALIGNMENT_MAP = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}

LINE_SPACING_MAP = {
    "single": WD_LINE_SPACING.SINGLE,
    "1.5": WD_LINE_SPACING.ONE_POINT_FIVE,
    "double": WD_LINE_SPACING.DOUBLE,
}

NUMBERING_STYLE_MAP = {
    "decimal": "decimal",
    "upper_letter": "upperLetter",
    "lower_letter": "lowerLetter",
    "upper_roman": "upperRoman",
    "lower_roman": "lowerRoman",
    "bullet": "bullet",
}

DEFAULT_FONT = {
    "name": "宋体",
    "size_pt": 12,
    "color": "#000000",
    "bold": False,
    "italic": False,
    "underline": False,
}

DEFAULT_PARAGRAPH = {
    "alignment": "left",
    "first_line_indent": None,
    "hanging_indent": None,
    "line_spacing": {"mode": "multiple", "value": 1.5},
    "space_before_pt": 0,
    "space_after_pt": 6,
}

DEFAULT_PAGE = {
    "paper_size": "A4",
    "orientation": "portrait",
    "margin_top_cm": 2.54,
    "margin_bottom_cm": 2.54,
    "margin_left_cm": 3.18,
    "margin_right_cm": 3.18,
}


def _parse_color(hex_color):
    """将 #RRGGBB 颜色转为 RGBColor"""
    if not hex_color or not hex_color.startswith("#"):
        return None
    hex_color = hex_color.lstrip("#")
    if len(hex_color) != 6:
        return None
    try:
        return RGBColor(int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16))
    except ValueError:
        return None


def _apply_font(run, font_spec, base_font=None):
    """对 Run 对象应用字体样式

    Args:
        run: python-docx Run 对象
        font_spec: 字体样式字典
        base_font: 基础字体样式（用于继承未指定的属性）
    """
    merged = {}
    if base_font:
        merged.update(base_font)
    if font_spec:
        merged.update(font_spec)

    if "name" in merged:
        run.font.name = merged["name"]
        r = run._element
        rPr = r.find(qn('w:rPr'))
        if rPr is None:
            rPr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
            r.insert(0, rPr)
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}></w:rFonts>')
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:eastAsia'), merged["name"])

    if "size_pt" in merged:
        run.font.size = Pt(merged["size_pt"])

    color = _parse_color(merged.get("color"))
    if color:
        run.font.color.rgb = color

    if merged.get("bold"):
        run.bold = True
    if merged.get("italic"):
        run.italic = True
    if merged.get("underline"):
        run.underline = True


def _apply_paragraph_format(paragraph, para_spec, base_para=None):
    """对 Paragraph 对象应用段落样式

    Args:
        paragraph: python-docx Paragraph 对象
        para_spec: 段落样式字典
        base_para: 基础段落样式
    """
    merged = {}
    if base_para:
        merged.update(base_para)
    if para_spec:
        merged.update(para_spec)

    if "alignment" in merged and merged["alignment"] in ALIGNMENT_MAP:
        paragraph.alignment = ALIGNMENT_MAP[merged["alignment"]]

    pf = paragraph.paragraph_format

    first_line = merged.get("first_line_indent")
    if first_line and isinstance(first_line, dict):
        value = first_line.get("value", 0)
        unit = first_line.get("unit", "char")
        if unit == "char":
            # 中文字符宽度 ≈ 字体大小
            font_size = (merged.get("font", {}) or {}).get("size_pt", 12) if isinstance(merged.get("font"), dict) else 12
            pf.first_line_indent = Cm(value * font_size * 0.035)
        elif unit == "cm":
            pf.first_line_indent = Cm(value)
        elif unit == "pt":
            pf.first_line_indent = Pt(value)

    hanging = merged.get("hanging_indent")
    if hanging and isinstance(hanging, dict):
        value = hanging.get("value", 0)
        unit = hanging.get("unit", "cm")
        if unit == "cm":
            pf.first_line_indent = Cm(-value)
        elif unit == "pt":
            pf.first_line_indent = Pt(-value)

    line_spacing = merged.get("line_spacing")
    if line_spacing and isinstance(line_spacing, dict):
        mode = line_spacing.get("mode", "multiple")
        value = line_spacing.get("value", 1.5)
        if mode == "multiple":
            pf.line_spacing = float(value)
        elif mode == "fixed":
            pf.line_spacing = Pt(float(value))
        elif mode in LINE_SPACING_MAP:
            pf.line_spacing = LINE_SPACING_MAP[mode]
    elif line_spacing and isinstance(line_spacing, (int, float)):
        pf.line_spacing = float(line_spacing)

    if "space_before_pt" in merged:
        pf.space_before = Pt(merged["space_before_pt"])
    if "space_after_pt" in merged:
        pf.space_after = Pt(merged["space_after_pt"])


def _setup_page(doc, page_spec):
    """设置页面属性"""
    merged = dict(DEFAULT_PAGE)
    if page_spec:
        merged.update(page_spec)

    section = doc.sections[0]

    paper = merged.get("paper_size", "A4")
    if paper in PAPER_SIZES:
        w, h = PAPER_SIZES[paper]
        if merged.get("orientation") == "landscape":
            section.orientation = WD_ORIENT.LANDSCAPE
            section.page_width = Cm(h)
            section.page_height = Cm(w)
        else:
            section.orientation = WD_ORIENT.PORTRAIT
            section.page_width = Cm(w)
            section.page_height = Cm(h)

    section.top_margin = Cm(merged.get("margin_top_cm", 2.54))
    section.bottom_margin = Cm(merged.get("margin_bottom_cm", 2.54))
    section.left_margin = Cm(merged.get("margin_left_cm", 3.18))
    section.right_margin = Cm(merged.get("margin_right_cm", 3.18))


def _setup_header_footer(doc, page_spec):
    """设置页眉页脚"""
    if not page_spec:
        return

    section = doc.sections[0]

    header_spec = page_spec.get("header")
    if header_spec and isinstance(header_spec, dict):
        header = section.header
        header.is_linked_to_previous = False
        text = header_spec.get("text", "")
        if text:
            p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            p.text = text
            align = header_spec.get("alignment", "center")
            if align in ALIGNMENT_MAP:
                p.alignment = ALIGNMENT_MAP[align]

    footer_spec = page_spec.get("footer")
    if footer_spec and isinstance(footer_spec, dict):
        footer = section.footer
        footer.is_linked_to_previous = False
        text = footer_spec.get("text", "")
        show_page_number = footer_spec.get("show_page_number", False)

        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        align = footer_spec.get("alignment", "center")
        if align in ALIGNMENT_MAP:
            p.alignment = ALIGNMENT_MAP[align]

        if text and show_page_number:
            p.text = text
            run = p.add_run()
            _add_page_number(run)
        elif show_page_number:
            _add_page_number(p.add_run())
        elif text:
            p.text = text


def _add_page_number(run):
    """添加页码域代码"""
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._r.append(fldChar1)
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    run._r.append(instrText)
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run._r.append(fldChar2)


def _add_heading(doc, section_data, default_font, default_para, warnings):
    """添加标题"""
    level = section_data.get("level", 1)
    content = section_data.get("content", "")
    heading = doc.add_heading(content, level=level)

    font_spec = section_data.get("font")
    para_spec = section_data.get("paragraph")

    if font_spec or para_spec:
        for run in heading.runs:
            _apply_font(run, font_spec, default_font)
        if para_spec:
            _apply_paragraph_format(heading, para_spec, default_para)


def _add_paragraph(doc, section_data, default_font, default_para, warnings):
    """添加段落，支持内联 runs（混合格式）"""
    runs = section_data.get("runs")
    content = section_data.get("content", "")

    p = doc.add_paragraph()
    font_spec = section_data.get("font")
    para_spec = section_data.get("paragraph")

    if runs and isinstance(runs, list):
        for run_spec in runs:
            text = run_spec.get("text", "")
            run = p.add_run(text)
            run_font = dict(font_spec) if font_spec else {}
            run_font.update({k: v for k, v in run_spec.items() if k != "text"})
            _apply_font(run, run_font, default_font)
    else:
        run = p.add_run(content)
        _apply_font(run, font_spec, default_font)

    _apply_paragraph_format(p, para_spec, default_para)


def _add_table(doc, section_data, default_font, default_para, warnings):
    """添加表格"""
    headers = section_data.get("headers", [])
    rows = section_data.get("rows", [])
    if not headers and not rows:
        warnings.append("表格缺少表头和数据行，已跳过")
        return

    num_cols = len(headers) if headers else (len(rows[0]) if rows else 1)
    num_rows = (1 if headers else 0) + len(rows)
    table = doc.add_table(rows=num_rows, cols=num_cols, style="Table Grid")

    col_widths = section_data.get("column_widths_cm")
    if col_widths:
        for i, w in enumerate(col_widths):
            if i < num_cols:
                for row in table.rows:
                    row.cells[i].width = Cm(w)

    header_style = section_data.get("header_style", {})
    cell_alignment = section_data.get("cell_alignment", "left")

    if headers:
        hdr_row = table.rows[0]
        for i, h in enumerate(headers):
            cell = hdr_row.cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(h))
            h_font = {
                "bold": header_style.get("bold", True),
                "size_pt": header_style.get("size_pt", default_font.get("size_pt", 12)),
            }
            h_color = header_style.get("font_color")
            if h_color:
                h_font["color"] = h_color
            _apply_font(run, h_font, default_font)
            if cell_alignment in ALIGNMENT_MAP:
                p.alignment = ALIGNMENT_MAP[cell_alignment]

            bg = header_style.get("bg_color")
            if bg:
                shading = parse_xml(
                    f'<w:shd {nsdecls("w")} w:fill="{bg.lstrip("#")}" w:val="clear"/>'
                )
                cell._tc.get_or_add_tcPr().append(shading)

    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + (1 if headers else 0)]
        for c_idx, cell_text in enumerate(row_data):
            if c_idx < num_cols:
                cell = row.cells[c_idx]
                cell.text = ""
                p = cell.paragraphs[0]
                run = p.add_run(str(cell_text))
                _apply_font(run, None, default_font)
                if cell_alignment in ALIGNMENT_MAP:
                    p.alignment = ALIGNMENT_MAP[cell_alignment]

    border_spec = section_data.get("border")
    if border_spec:
        _apply_table_border(table, border_spec)

    doc.add_paragraph()


def _apply_table_border(table, border_spec):
    """设置表格边框"""
    style = border_spec.get("style", "single")
    size = border_spec.get("size", 4)
    color = border_spec.get("color", "000000").lstrip("#")

    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}></w:tblPr>')

    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="{style}" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'<w:left w:val="{style}" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'<w:bottom w:val="{style}" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'<w:right w:val="{style}" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'<w:insideH w:val="{style}" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'<w:insideV w:val="{style}" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)


def _add_list(doc, section_data, default_font, default_para, warnings):
    """添加列表"""
    items = section_data.get("items", [])
    ordered = section_data.get("ordered", False)
    style = section_data.get("style", "decimal" if ordered else "bullet")

    if not items:
        return

    for idx, item in enumerate(items):
        p = doc.add_paragraph()
        if ordered:
            prefix_map = {
                "decimal": f"{idx + 1}.",
                "upper_letter": f"{chr(65 + min(idx, 25))}.",
                "lower_letter": f"{chr(97 + min(idx, 25))}.",
                "upper_roman": f"{_to_roman(idx + 1)}.",
                "lower_roman": f"{_to_roman(idx + 1).lower()}.",
            }
            prefix = prefix_map.get(style, f"{idx + 1}.")
            run = p.add_run(f"{prefix} {item}")
        else:
            bullet_map = {
                "bullet": "\u2022",
                "circle": "\u25CB",
                "square": "\u25A0",
                "dash": "\u2013",
            }
            bullet = bullet_map.get(style, "\u2022")
            run = p.add_run(f"{bullet} {item}")

        _apply_font(run, section_data.get("font"), default_font)
        _apply_paragraph_format(p, section_data.get("paragraph"), default_para)


def _to_roman(n):
    """整数转罗马数字"""
    val = [1000, 900, 500, 400, 100, 90, 50, 40, 10, 9, 5, 4, 1]
    sym = ["M", "CM", "D", "CD", "C", "XC", "L", "XL", "X", "IX", "V", "IV", "I"]
    roman = ""
    for i in range(len(val)):
        while n >= val[i]:
            roman += sym[i]
            n -= val[i]
    return roman


def _add_image(doc, section_data, default_para, warnings):
    """添加图片"""
    path = section_data.get("path", "")
    if not path or not os.path.exists(path):
        warnings.append(f"图片文件不存在: {path}")
        return

    p = doc.add_paragraph()
    width = section_data.get("width_cm")
    height = section_data.get("height_cm")
    alignment = section_data.get("alignment", "center")

    if alignment in ALIGNMENT_MAP:
        p.alignment = ALIGNMENT_MAP[alignment]

    run = p.add_run()
    try:
        if width and height:
            run.add_picture(path, width=Cm(width), height=Cm(height))
        elif width:
            run.add_picture(path, width=Cm(width))
        elif height:
            run.add_picture(path, height=Cm(height))
        else:
            run.add_picture(path, width=Cm(10))
    except Exception as e:
        warnings.append(f"插入图片失败: {str(e)}")


def create_word_document(filename, content=None, formatting=None, output_dir=None):
    """创建格式化 Word 文档

    Args:
        filename: 文件名（不含扩展名）
        content: 纯文本内容（formatting 为 None 时使用，向后兼容）
        formatting: 排版描述 JSON 对象

    Returns:
        str: 结果消息，包含文件路径和可能的警告
    """
    if output_dir is None:
        output_dir = os.path.join("document_output", "word_output")
    os.makedirs(output_dir, exist_ok=True)
    filepath = os.path.join(output_dir, f"{filename}.docx")

    warnings = []
    doc = Document()

    if formatting and isinstance(formatting, dict):
        page_spec = formatting.get("page", {})
        default_font = dict(DEFAULT_FONT)
        if formatting.get("default_font"):
            default_font.update(formatting["default_font"])
        default_para = dict(DEFAULT_PARAGRAPH)
        if formatting.get("default_paragraph"):
            default_para.update(formatting["default_paragraph"])

        _setup_page(doc, page_spec)
        _setup_header_footer(doc, page_spec)

        sections = formatting.get("sections", [])
        if not sections:
            warnings.append("未指定内容段落(sections)，生成了空白文档")

        for sec in sections:
            sec_type = sec.get("type", "paragraph")
            if sec_type == "heading":
                _add_heading(doc, sec, default_font, default_para, warnings)
            elif sec_type == "paragraph":
                _add_paragraph(doc, sec, default_font, default_para, warnings)
            elif sec_type == "table":
                _add_table(doc, sec, default_font, default_para, warnings)
            elif sec_type == "list":
                _add_list(doc, sec, default_font, default_para, warnings)
            elif sec_type == "image":
                _add_image(doc, sec, default_para, warnings)
            elif sec_type == "page_break":
                doc.add_page_break()
            else:
                warnings.append(f"未知的段落类型 '{sec_type}'，已跳过")
    else:
        doc.add_heading(filename, level=1)
        text = content or ""
        for para in [p.strip() for p in text.replace("\r", "").split("\n") if p.strip()]:
            doc.add_paragraph(para)

    doc.save(filepath)

    result = f"Word文档已成功保存至: {filepath}"
    if warnings:
        result += "\n[排版提示] " + "; ".join(warnings)
    return result